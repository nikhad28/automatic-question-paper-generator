import os
import json
import re
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
import openai
import PyPDF2
import docx
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
import hashlib
import traceback

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'dev_secret_key_123')
CORS(app)

# Configure upload folder
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# ======================
# OPENAI API INITIALIZATION
# ======================

OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
openai_client = None

if OPENAI_API_KEY and OPENAI_API_KEY != 'your_openai_api_key_here':
    openai_client = openai.OpenAI(api_key=OPENAI_API_KEY)
    print("✅ OpenAI API initialized")
else:
    print("⚠️ OpenAI API key not set. Please add OPENAI_API_KEY to .env file.")

# ======================
# USER STORAGE
# ======================

users_storage = {}

# ======================
# HELPER FUNCTIONS
# ======================

def extract_text_from_file(file_path):
    try:
        file_ext = file_path.split('.')[-1].lower()
        
        if file_ext == 'pdf':
            return extract_text_from_pdf(file_path)
        elif file_ext == 'docx':
            return extract_text_from_docx(file_path)
        elif file_ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return ""
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return text

def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"DOCX extraction error: {e}")
    return text

def generate_questions_with_openai(syllabus_text, paper_type, num_questions=10, marks_per_question=2):
    """Generate questions using OpenAI API"""
    
    if not openai_client:
        # Fallback to OpenRouter or fallback
        return generate_questions_with_ai_fallback(syllabus_text, paper_type, num_questions, marks_per_question)
    
    try:
        # Determine question type
        type_desc = {
            'mcq': f"{num_questions} Multiple Choice Questions with 4 options each. Each question must have proper options (A, B, C, D) with meaningful choices.",
            'descriptive': f"{num_questions} Descriptive questions. Include What, Why, Explain, Differentiate type questions.",
            'ut': "Unit Test paper (20 marks total). Mix of short (2 marks), medium (4 marks), long (6 marks) questions.",
            'semester': "Semester Exam paper (80 marks total). Mix of 1, 2, 5, 10 mark questions."
        }.get(paper_type, f"{num_questions} questions")
        
        if paper_type == 'ut':
            total_marks = 20
        elif paper_type == 'semester':
            total_marks = 80
        else:
            total_marks = num_questions * marks_per_question
        
        # Build system prompt
        system_prompt = """You are an expert university professor creating high-quality question papers. 
        You generate questions based ONLY on the provided syllabus content.
        Each question must be mapped to a Course Outcome (CO1-CO6).
        Use proper academic language and university-level difficulty.
        Return ONLY valid JSON format with no additional text."""
        
        # Build user prompt
        user_prompt = f"""
SYLLABUS:
{syllabus_text[:6000]}

TASK: Generate {type_desc}

OUTPUT FORMAT (JSON only):
{{
    "title": "University Question Paper",
    "course_code": "Based on Syllabus",
    "total_marks": {total_marks},
    "time_duration": "3 Hours",
    "instructions": "All questions are compulsory. Attempt all questions.",
    "questions": [
        {{
            "id": 1,
            "type": "mcq/descriptive/short_answer/long_answer",
            "question": "Clear question text",
            "options": ["A. Option 1", "B. Option 2", "C. Option 3", "D. Option 4"],
            "answer": "A. Option 1",
            "marks": 2,
            "co": "CO1",
            "blooms_level": "Remember/Understand/Apply/Analyze/Evaluate/Create"
        }}
    ]
}}

Generate exactly {num_questions} questions. Return ONLY valid JSON."""
        
        print(f"🔄 Generating {num_questions} questions using OpenAI...")
        
        # Call OpenAI API
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",  # or "gpt-3.5-turbo" for cheaper option
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=4096
        )
        
        result_text = response.choices[0].message.content
        print("📝 Parsing OpenAI response...")
        
        # Parse JSON
        try:
            json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
            if json_match:
                questions_data = json.loads(json_match.group())
            else:
                questions_data = json.loads(result_text)
            
            if 'questions' not in questions_data:
                questions_data['questions'] = []
            
            for idx, q in enumerate(questions_data.get('questions', [])):
                if not q.get('id'):
                    q['id'] = idx + 1
                if not q.get('co'):
                    q['co'] = f"CO{(idx % 6) + 1}"
                if q.get('type') == 'mcq' and q.get('options'):
                    # Ensure options are properly formatted
                    if len(q['options']) == 4:
                        pass  # Already correct
                    elif isinstance(q['options'], list) and len(q['options']) > 0:
                        # Reformat options if needed
                        pass
            
            print(f"✅ Generated {len(questions_data['questions'])} questions using OpenAI")
            
            return {
                "success": True,
                "data": questions_data,
                "ai_generated": True,
                "model": "OpenAI GPT-4"
            }
            
        except json.JSONDecodeError as e:
            print(f"❌ JSON Parse Error: {e}")
            print(f"Response preview: {result_text[:500]}")
            return generate_questions_with_ai_fallback(syllabus_text, paper_type, num_questions, marks_per_question)
            
    except Exception as e:
        print(f"❌ OpenAI Error: {e}")
        traceback.print_exc()
        return generate_questions_with_ai_fallback(syllabus_text, paper_type, num_questions, marks_per_question)

def generate_questions_with_ai_fallback(syllabus_text, paper_type, num_questions=10, marks_per_question=2):
    """Fallback: Use OpenRouter or generate sample questions"""
    
    # Try OpenRouter first if available
    try:
        import requests
        OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY')
        if OPENROUTER_API_KEY:
            # Try OpenRouter with Gemini
            headers = {
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": "http://localhost:5000",
                "X-Title": "Question Paper Generator"
            }
            
            type_desc = {
                'mcq': f"{num_questions} Multiple Choice Questions with 4 options each. Each question must have proper options (A, B, C, D) with meaningful choices.",
                'descriptive': f"{num_questions} Descriptive questions.",
                'ut': "Unit Test paper (20 marks total)",
                'semester': "Semester Exam paper (80 marks total)"
            }.get(paper_type, f"{num_questions} questions")
            
            if paper_type == 'ut':
                total_marks = 20
            elif paper_type == 'semester':
                total_marks = 80
            else:
                total_marks = num_questions * marks_per_question
            
            prompt = f"""You are an expert university professor. Generate a question paper based on this syllabus:

SYLLABUS:
{syllabus_text[:4000]}

TASK: Generate {type_desc}

OUTPUT FORMAT (JSON only):
{{
    "title": "Question Paper",
    "total_marks": {total_marks},
    "questions": [
        {{
            "id": 1,
            "type": "mcq/descriptive",
            "question": "Question text",
            "options": ["A. Option 1", "B. Option 2", "C. Option 3", "D. Option 4"],
            "answer": "A. Option 1",
            "marks": 2,
            "co": "CO1"
        }}
    ]
}}

Generate {num_questions} questions. Return ONLY valid JSON."""
            
            data = {
                "model": "google/gemini-1.5-flash",
                "messages": [
                    {"role": "system", "content": "You are an expert university professor. Return ONLY valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 4096
            }
            
            response = requests.post("https://openrouter.ai/api/v1/chat/completions", 
                                     headers=headers, json=data, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    questions_data = json.loads(json_match.group())
                else:
                    questions_data = json.loads(content)
                
                if 'questions' not in questions_data:
                    questions_data['questions'] = []
                
                for idx, q in enumerate(questions_data.get('questions', [])):
                    if not q.get('id'):
                        q['id'] = idx + 1
                    if not q.get('co'):
                        q['co'] = f"CO{(idx % 6) + 1}"
                
                print(f"✅ Generated {len(questions_data['questions'])} questions using OpenRouter")
                
                return {
                    "success": True,
                    "data": questions_data,
                    "ai_generated": True,
                    "model": "OpenRouter Gemini"
                }
    except Exception as e:
        print(f"OpenRouter fallback error: {e}")
    
    # Ultimate fallback - sample questions
    print("📝 Using fallback sample questions...")
    return generate_sample_questions(syllabus_text, paper_type, num_questions, marks_per_question)

def generate_sample_questions(syllabus_text, paper_type, num_questions=10, marks_per_question=2):
    """Generate sample questions as ultimate fallback"""
    questions = []
    co_list = ['CO1', 'CO2', 'CO3', 'CO4', 'CO5', 'CO6']
    blooms_levels = ['Remember', 'Understand', 'Apply', 'Analyze', 'Evaluate', 'Create']
    
    # Extract key terms
    import re
    words = re.findall(r'\b[A-Za-z]{4,}\b', syllabus_text)
    stopwords = {'this', 'that', 'these', 'those', 'with', 'from', 'have', 'will', 'can', 'should', 'would', 'could', 'then', 'than', 'they', 'them', 'their', 'there', 'what', 'when', 'where', 'which', 'while'}
    key_terms = []
    for word in words:
        if word.lower() not in stopwords and len(word) > 3 and word not in key_terms:
            key_terms.append(word)
    if not key_terms:
        key_terms = ["concepts", "principles", "applications", "theories", "methodologies"]
    
    if paper_type == 'mcq':
        templates = [
            {"q": "What is the primary purpose of {term}?", "opts": ["A. To analyze and process {term} effectively", "B. To store and retrieve {term} efficiently", "C. To implement {term} in applications", "D. To optimize {term} for performance"]},
            {"q": "Which of the following best describes {term}?", "opts": ["A. A systematic approach to problem-solving", "B. A set of rules for data processing", "C. A framework for building applications", "D. A methodology for system design"]},
            {"q": "What is the main function of {term}?", "opts": ["A. To manage and organize data", "B. To execute computational tasks", "C. To process information efficiently", "D. To support decision-making"]},
            {"q": "What does {term} primarily focus on?", "opts": ["A. Problem-solving and optimization", "B. Data processing and analysis", "C. System design and implementation", "D. Algorithm development"]}
        ]
        for i in range(min(num_questions, 15)):
            term = key_terms[i % len(key_terms)]
            tmpl = templates[i % len(templates)]
            questions.append({
                "id": i + 1,
                "type": "mcq",
                "question": tmpl["q"].format(term=term),
                "options": [opt.format(term=term) for opt in tmpl["opts"]],
                "answer": tmpl["opts"][0].format(term=term),
                "marks": marks_per_question,
                "co": co_list[i % 6],
                "blooms_level": blooms_levels[i % 6]
            })
    else:
        templates = [
            "Explain the concept of {term} with relevant examples.",
            "Discuss the importance of {term} in detail.",
            "Describe the key features and applications of {term}.",
            "Analyze the role of {term} in modern context.",
            "Evaluate the significance of {term}."
        ]
        for i in range(min(num_questions, 15)):
            term = key_terms[i % len(key_terms)]
            template = templates[i % len(templates)]
            questions.append({
                "id": i + 1,
                "type": "descriptive",
                "question": template.format(term=term),
                "marks": marks_per_question,
                "co": co_list[i % 6],
                "blooms_level": blooms_levels[i % 6]
            })
    
    total_marks = sum(q['marks'] for q in questions)
    
    return {
        "success": True,
        "data": {
            "title": f"{paper_type.upper()} Question Paper",
            "total_marks": total_marks,
            "questions": questions
        },
        "ai_generated": False,
        "note": "Sample questions generated (fallback mode)"
    }

def generate_pdf_paper(questions_data, filename):
    try:
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=24, textColor=colors.darkred, alignment=TA_CENTER, spaceAfter=30)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'], fontSize=16, textColor=colors.darkblue, spaceAfter=12)
        question_style = ParagraphStyle('QuestionStyle', parent=styles['Normal'], fontSize=12, spaceAfter=8)
        
        content = []
        title = questions_data.get('title', 'Question Paper')
        content.append(Paragraph(title, title_style))
        content.append(Spacer(1, 12))
        course_info = f"Total Marks: {questions_data.get('total_marks', 0)}"
        content.append(Paragraph(course_info, heading_style))
        content.append(Spacer(1, 20))
        
        for q in questions_data.get('questions', []):
            q_text = f"<b>Q{q['id']}.</b> {q['question']}"
            if q.get('type') == 'mcq' and q.get('options'):
                q_text += "<br/>"
                for opt in q.get('options', []):
                    q_text += f"&nbsp;&nbsp;&nbsp;• {opt}<br/>"
            q_text += f"<br/><i>[Marks: {q.get('marks', 1)}, CO: {q.get('co', 'N/A')}]</i>"
            content.append(Paragraph(q_text, question_style))
            content.append(Spacer(1, 8))
        
        content.append(Spacer(1, 20))
        content.append(Paragraph("--- End of Paper ---", styles['Normal']))
        doc.build(content)
        return True
    except Exception as e:
        print(f"PDF generation error: {e}")
        return False

# ======================
# ROUTES
# ======================

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        if email == 'admin@test.com' and password == 'admin123':
            session['user_id'] = 'admin'
            session['user_name'] = 'Admin'
            return redirect(url_for('dashboard'))
        
        if email in users_storage:
            user_data = users_storage[email]
            if user_data.get('password') == hashlib.sha256(password.encode()).hexdigest():
                session['user_id'] = email
                session['user_name'] = user_data.get('name', email)
                return redirect(url_for('dashboard'))
        
        error = "Invalid credentials"
        return render_template('login.html', error=error)
    
    return render_template('login.html')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    error = None
    if request.method == 'POST':
        name = request.form.get('name')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        if password != confirm_password:
            error = "Passwords do not match"
            return render_template('signup.html', error=error)
        
        if email in users_storage:
            error = "Email already registered"
            return render_template('signup.html', error=error)
        
        users_storage[email] = {
            'name': name,
            'email': email,
            'password': hashlib.sha256(password.encode()).hexdigest()
        }
        return redirect(url_for('login'))
    
    return render_template('signup.html')

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return render_template('dashboard.html', name=session.get('user_name', 'User'))

@app.route('/generate/<paper_type>')
def generate_page(paper_type):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    titles = {
        'mcq': 'MCQ Question Generator',
        'descriptive': 'Descriptive Questions Generator',
        'ut': 'Unit Test Paper Generator (20 Marks)',
        'semester': 'Semester Paper Generator (80 Marks)'
    }
    
    return render_template('generate.html', 
                         paper_type=paper_type,
                         title=titles.get(paper_type, 'Question Generator'),
                         name=session.get('user_name', 'User'))

# ======================
# API ENDPOINTS
# ======================

@app.route('/api/generate-paper', methods=['POST'])
def generate_paper():
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    
    try:
        paper_type = request.form.get('paper_type')
        syllabus_text = request.form.get('syllabus_text', '')
        num_questions = int(request.form.get('num_questions', 10))
        marks_per_question = int(request.form.get('marks_per_question', 2))
        
        file = request.files.get('syllabus_file')
        if file and file.filename:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)
            extracted_text = extract_text_from_file(file_path)
            if extracted_text:
                syllabus_text += "\n" + extracted_text
            os.remove(file_path)
        
        if not syllabus_text or len(syllabus_text.strip()) < 10:
            return jsonify({
                'success': False,
                'error': 'Please provide proper syllabus content (minimum 10 characters)'
            }), 400
        
        # Try OpenAI first
        result = generate_questions_with_openai(syllabus_text, paper_type, num_questions, marks_per_question)
        
        if result['success']:
            session['current_paper'] = result['data']
            return jsonify({
                'success': True,
                'data': result['data'],
                'ai_generated': result.get('ai_generated', False),
                'model': result.get('model', 'unknown')
            })
        else:
            return jsonify({
                'success': False,
                'error': result.get('error', 'Failed to generate questions')
            }), 500
            
    except Exception as e:
        print(f"Generate paper error: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/download-paper/<format>')
def download_paper(format):
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    
    paper_data = session.get('current_paper')
    if not paper_data:
        return jsonify({'success': False, 'error': 'No paper data found. Please generate a paper first.'}), 400
    
    try:
        if format == 'pdf':
            filename = f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            if generate_pdf_paper(paper_data, temp_path):
                return send_file(temp_path, as_attachment=True, download_name=filename)
            else:
                return jsonify({'success': False, 'error': 'PDF generation failed'}), 500
                
        elif format == 'docx':
            doc = docx.Document()
            
            title = doc.add_heading(paper_data.get('title', 'Question Paper'), 0)
            title.alignment = 1
            
            doc.add_paragraph(f"Total Marks: {paper_data.get('total_marks', 0)}")
            doc.add_paragraph()
            
            for q in paper_data.get('questions', []):
                p = doc.add_paragraph()
                p.add_run(f"Q{q['id']}.").bold = True
                p.add_run(f" {q['question']}")
                
                if q.get('type') == 'mcq' and q.get('options'):
                    for opt in q.get('options', []):
                        doc.add_paragraph(f"    • {opt}")
                
                doc.add_paragraph(f"[Marks: {q.get('marks', 1)}, CO: {q.get('co', 'N/A')}]")
                doc.add_paragraph()
            
            filename = f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            doc.save(temp_path)
            
            return send_file(temp_path, as_attachment=True, download_name=filename)
            
        else:
            return jsonify({'success': False, 'error': 'Invalid format. Use pdf or docx'}), 400
            
    except Exception as e:
        print(f"Download error: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))

# ======================
# MAIN
# ======================

if __name__ == '__main__':
    print("\n" + "="*50)
    print("📝 AI Question Paper Generator")
    print("="*50)
    print(f"🚀 Server running on: http://127.0.0.1:5000")
    print(f"🤖 OpenAI API: {'✅ Connected' if openai_client else '❌ Not connected'}")
    print("="*50 + "\n")
    
    app.run(debug=True, host='0.0.0.0', port=5000)